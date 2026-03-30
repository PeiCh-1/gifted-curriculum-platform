from docx import Document
import sys

def inspect_doc(path):
    print(f"\n--- Inspecting {path} ---")
    try:
        doc = Document(path)
    except Exception as e:
        print(f"Failed to open {path}: {e}")
        return
        
    print("Paragraphs:")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text: print(f"[{i}] {text}")
        
    print("\nTables:")
    for i, t in enumerate(doc.tables):
        print(f"Table {i}: {len(t.rows)} rows, {len(t.columns)} cols")
        for r_idx, r in enumerate(t.rows):
            row_text = []
            for c in r.cells:
                row_text.append(c.text.replace("\n", " ").strip())
            print(f"  Row {r_idx}: | " + " | ".join(row_text) + " |")

if __name__ == '__main__':
    inspect_doc('public/curriculum_template.docx')
    inspect_doc('public/igp_template.docx')
