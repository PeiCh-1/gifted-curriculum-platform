import sys
import os
import shutil
from docx import Document
from docx.shared import RGBColor

# 原始範本路徑 (包含可能的空格)
SRC_CURRICULUM = '../課程計劃草稿/要繳交的資料格式/附件3特殊教育課程計畫格式＿空白格式 .docx'
SRC_IGP = '../課程計劃草稿/要繳交的資料格式/IGP個別調整表.docx'

def patch_curriculum(path):
    if not os.path.exists(SRC_CURRICULUM):
        print(f"Error: Original curriculum source not found at {SRC_CURRICULUM}")
        return
    
    # 每次都先還原為最乾淨的版本
    shutil.copy(SRC_CURRICULUM, path)
    doc = Document(path)
    
    # 1. 處理填表教師 (全域搜尋段落與表格)
    found_teacher = False
    for p in doc.paragraphs:
        if '填表教師' in p.text and '{Teacher}' not in p.text:
            p.add_run(' {Teacher}')
            found_teacher = True
    
    # 2. 更新表格標籤
    for t in doc.tables:
        if len(t.rows) >= 6:
            # 填寫基本資料
            row0 = t.rows[0]
            for i, c in enumerate(row0.cells):
                if "領域/科目" in c.text and i+1 < len(row0.cells):
                    row0.cells[i+1].text = "{DomainModeString}"
                elif "課程名稱" in c.text and i+1 < len(row0.cells):
                    row0.cells[i+1].text = "{CourseName}"
            
            row1 = t.rows[1]
            for i, c in enumerate(row1.cells):
                if "年級/組別" in c.text and i+1 < len(row1.cells):
                    row1.cells[i+1].text = "{Grade}"
                elif "教材來源" in c.text and i+1 < len(row1.cells):
                    row1.cells[i+1].text = "{MaterialSource}"

            row2 = t.rows[2]
            row2.cells[1].text = "{WeeklyPeriods}"
            row2.cells[-1].text = "{Teacher}" # 表格內的教學者

            row3 = t.rows[3]
            row3.cells[1].text = "{CoreCompetencies}"

            # 關鍵：富文本樣式迴圈 (Indicators 欄位)
            row5 = t.rows[5]
            row5.cells[0].text = "{#Weeks}{Week}"
            
            # --- 注入富文本標籤 ---
            cell = row5.cells[1]
            cell.text = "" # 清空
            p = cell.paragraphs[0]
            p.add_run("{#IndRuns}{#isAdd}")
            r_add = p.add_run("{text}")
            r_add.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # 紅字
            p.add_run("{/isAdd}{#isDel}")
            r_del = p.add_run("{text}")
            r_del.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # 紅字
            r_del.font.strike = True # 刪除線
            p.add_run("{/isDel}{^isAdd}{^isDel}{text}{/isDel}{/isAdd}{/IndRuns}")
            # --- 結束注入 ---

            row5.cells[2].text = "{LessonFocus}"
            row5.cells[3].text = "{Assessment}"
            row5.cells[5].text = "{Issues}"
            row5.cells[7].text = "{Notes}{/Weeks}"
            
            # 刪除多餘 Row
            for i in range(len(t.rows)-1, 5, -1):
                tr = t.rows[i]._tr
                tr.getparent().remove(tr)

    # 3. 更新備註 (清理末尾段落並重新注入)
    # 移除段落中包含重複項的內容
    for p in doc.paragraphs:
        if "實際上課日數" in p.text or "融入議題參考" in p.text:
            p.text = "" # 清空

    notes_parts = [
        "1.本學年實際上課日數及補休補班調整，仍依教育局公告之本學年度重要行事曆辦理。",
        "2.融入議題參考：性別平等教育、人權教育、環境教育、海洋教育、科技教育、能源教育、家庭教育、原住民族教育、品德教育、生命教育、法治教育、資訊教育、安全教育、防災教育、生涯規劃教育、多元文化教育、閱讀素養教育、戶外教育、國際教育…等（上述議題係參考「十二年國教課綱議題融入說明手冊」所列出，各校亦可選擇適合之議題填入）。",
        "3.評量方式填寫參考：口頭評量、紙筆評量、實作評量、教師觀察、學生自評、同儕互評或其他適合之評量方式。"
    ]
    for note in notes_parts:
        doc.add_paragraph(note)

    doc.save(path)
    print(f"Patched Curriculum Template: {path}")

def patch_igp(path):
    if not os.path.exists(SRC_IGP):
        print(f"Error: Original IGP source not found at {SRC_IGP}")
        return
    
    shutil.copy(SRC_IGP, path)
    doc = Document(path)
    
    for t in doc.tables:
        if len(t.rows) >= 3:
            # Row 0: Header
            row0 = t.rows[0]
            row0.cells[1].text = "{CourseName}"
            row0.cells[3].text = "{CourseType}"
            row0.cells[5].text = "{Teacher}"
            
            # Row 2: Indicators with Rich Text Loop
            row2 = t.rows[2]
            row2.cells[0].text = "1"
            
            cell = row2.cells[1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.add_run("{#IndRuns}{#isAdd}")
            r_add = p.add_run("{text}")
            r_add.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.add_run("{/isAdd}{#isDel}")
            r_del = p.add_run("{text}")
            r_del.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r_del.font.strike = True
            p.add_run("{/isDel}{^isAdd}{^isDel}{text}{/isDel}{/isAdd}{/IndRuns}")
            
            row2.cells[-1].text = "{GlobalStrategies}"
            
            # 刪除多餘 Row
            for i in range(len(t.rows)-1, 2, -1):
                tr = t.rows[i]._tr
                tr.getparent().remove(tr)

    doc.save(path)
    print(f"Patched IGP Template: {path}")

if __name__ == "__main__":
    patch_curriculum('public/curriculum_template.docx')
    patch_igp('public/igp_template.docx')
